"""RAG tools for document ingestion, retrieval, and response generation.

Phase 3.1: ChromaDB implementation with async file I/O.
Supports ".pdf", ".md", ".txt", ".docx",
    ".py", ".js", ".ts", ".jsx", ".tsx", ".json", ".rst"
"""

import asyncio
import logging
from pathlib import Path
from typing import TYPE_CHECKING, List, Optional

import aiofiles
from langchain_community.embeddings import OllamaEmbeddings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from docx import Document as DocxDocument

from src.core.config import settings
from src.core.model_router import model_router

from langchain_chroma import Chroma


logger = logging.getLogger(__name__)

# Supported file extensions
SUPPORTED_EXTENSIONS = {
    ".pdf", ".md", ".txt", ".docx",
    ".py", ".js", ".ts", ".jsx", ".tsx", ".json", ".rst"
}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def get_embedding_model(model_name: str) -> OllamaEmbeddings:
    """Create an OllamaEmbeddings instance for the specified model.

    Args:
        model_name: Name of the embedding model (e.g., "nomic-embed-text", "bge-m3")

    Returns:
        OllamaEmbeddings instance configured for the model

    Raises:
        ValueError: If model_name is invalid or empty
    """
    if not model_name or not isinstance(model_name, str):
        raise ValueError(f"Invalid model name: {model_name}")

    logger.info(
        f"Initializing embedding model: {model_name}",
        extra={"model": model_name, "ollama_host": settings.OLLAMA_HOST},
    )

    try:
        embeddings = OllamaEmbeddings(
            model=model_name,
            base_url=settings.OLLAMA_HOST,
        )
        logger.debug(f"Embedding model '{model_name}' initialized successfully")
        return embeddings
    except Exception as e:
        logger.error(
            f"Failed to initialize embedding model '{model_name}': {e}",
            extra={"model": model_name, "error": str(e)},
            exc_info=True,
        )
        raise


def get_vector_store(
    backend: str,
    embed_model: OllamaEmbeddings,
    collection_name: str,
) -> "Chroma":  # type: ignore
    """Create a vector store instance based on the specified backend.

    Supports ChromaDB (local) and Qdrant Cloud (remote).

    Args:
        backend: Vector store backend ("chroma" or "qdrant")
        embed_model: OllamaEmbeddings instance for generating embeddings
        collection_name: Name of the collection to use or create

    Returns:
        LangChain-compatible vector store (Chroma or QdrantVectorStore instance)

    Raises:
        ValueError: If backend is not supported or required config is missing
        Exception: If vector store initialization fails
    """
    if backend not in ["chroma", "qdrant"]:
        raise ValueError(f"Unsupported vector backend: {backend}. Supported: chroma, qdrant")

    logger.info(
        f"Initializing vector store: {backend}",
        extra={"backend": backend, "collection": collection_name},
    )

    try:
        if backend == "chroma":
            from langchain_chroma import Chroma
            import chromadb

            # Ensure persist directory exists
            persist_dir = Path(settings.CHROMA_PERSIST_DIR)
            persist_dir.mkdir(parents=True, exist_ok=True)

            # Create persistent ChromaDB client
            chroma_client = chromadb.PersistentClient(path=str(persist_dir))

            # Create Chroma vector store with LangChain wrapper
            vector_store = Chroma(
                client=chroma_client,
                collection_name=collection_name,
                embedding_function=embed_model,
            )

            logger.info(
                f"ChromaDB vector store initialized: collection='{collection_name}', persist_dir='{persist_dir}'",
                extra={
                    "backend": "chroma",
                    "collection": collection_name,
                    "persist_dir": str(persist_dir),
                },
            )
            return vector_store

        elif backend == "qdrant":
            from langchain_qdrant import QdrantVectorStore
            from qdrant_client import QdrantClient

            # Validate required Qdrant configuration
            if not settings.QDRANT_URL:
                raise ValueError("QDRANT_URL is required for Qdrant backend. Set it in .env")
            if not settings.QDRANT_API_KEY:
                raise ValueError("QDRANT_API_KEY is required for Qdrant backend. Set it in .env")

            # Create Qdrant client with timeout for cloud requests
            qdrant_client = QdrantClient(
                url=settings.QDRANT_URL,
                api_key=settings.QDRANT_API_KEY,
                timeout=30,  # 30 second timeout for cloud requests
            )

            # Create Qdrant vector store with LangChain wrapper
            vector_store = QdrantVectorStore(
                client=qdrant_client,
                collection_name=collection_name or settings.QDRANT_COLLECTION,
                embedding=embed_model,
            )

            logger.info(
                f"Qdrant Cloud vector store initialized: collection='{collection_name}', url='{settings.QDRANT_URL}'",
                extra={
                    "backend": "qdrant",
                    "collection": collection_name,
                    "url": settings.QDRANT_URL,
                },
            )
            return vector_store

    except Exception as e:
        logger.error(
            f"Failed to initialize vector store '{backend}': {e}",
            extra={"backend": backend, "error": str(e)},
            exc_info=True,
        )
        raise


async def read_document(file_path: str) -> str:
    """Read a document file asynchronously and extract text content.

    Supports PDF, Markdown, TXT, DOCX, and Code formats.

    Args:
        file_path: Path to the document file

    Returns:
        Extracted text content as string

    Raises:
        ValueError: If file path is invalid, extension not supported, or file too large
        FileNotFoundError: If file does not exist
        Exception: If file reading fails
    """
    # Normalize and validate file path
    path = Path(file_path).resolve()

    # Security: Prevent directory traversal
    if ".." in str(path):
        raise ValueError(f"Invalid file path (directory traversal detected): {file_path}")

    # Check file exists
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    # Check file extension
    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: {extension}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    # Check file size
    file_size = path.stat().st_size
    if file_size > MAX_FILE_SIZE:
        raise ValueError(
            f"File too large: {file_size} bytes (max: {MAX_FILE_SIZE} bytes). File: {file_path}"
        )

    logger.info(
        f"Reading document: {file_path}",
        extra={"file_path": str(path), "extension": extension, "size": file_size},
    )

    try:
        if extension == ".pdf":
            # Read PDF using pypdf (synchronous library, run in thread pool)
            def read_pdf():
                reader = PdfReader(str(path))
                text = "\n".join([page.extract_text() for page in reader.pages])
                return text

            text = await asyncio.to_thread(read_pdf)

        elif extension in {".md", ".txt", ".py", ".js", ".ts", ".jsx", ".tsx", ".json", ".rst"}:
            # Read text and code files asynchronously
            async with aiofiles.open(path, mode="r", encoding="utf-8") as f:
                text = await f.read()

        elif extension == ".docx":
            # Read DOCX using python-docx (synchronous library, run in thread pool)
            def read_docx():
                doc = DocxDocument(str(path))
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return text

            text = await asyncio.to_thread(read_docx)

        else:
            # Should not reach here due to extension check above
            raise ValueError(f"Unsupported file type: {extension}")

        logger.debug(
            f"Document read successfully: {file_path} ({len(text)} characters)",
            extra={"file_path": str(path), "char_count": len(text)},
        )

        return text

    except Exception as e:
        logger.error(
            f"Failed to read document '{file_path}': {e}",
            extra={"file_path": str(path), "error": str(e)},
            exc_info=True,
        )
        raise


def chunk_document(
    text: str,
    file_path: str,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
) -> List[Document]:
    """
    Split a document using code-aware or text chunking.
    
    Phase 10.1: Uses Tree-sitter for code, falls back to text.
    """
    logger.debug(
        f"Chunking document: {file_path}",
        extra={"file_path": file_path, "chunk_size": chunk_size, "chunk_overlap": chunk_overlap},
    )

    try:
        # Phase 10.1: Try code-aware chunking
        try:
            from src.agents.rag.chunking import CodeChunker, TextChunker
            
            code_chunker = CodeChunker()
            if code_chunker.is_supported(file_path):
                chunks_data = code_chunker.chunk(text, file_path)
            else:
                text_chunker = TextChunker(chunk_size, chunk_overlap)
                chunks_data = text_chunker.chunk(text, file_path)
            
            # Convert to LangChain Document format
            documents = []
            for chunk_data in chunks_data:
                doc = Document(
                    page_content=chunk_data["content"],
                    metadata=chunk_data["metadata"]
                )
                documents.append(doc)
            
            logger.debug(f"Code-aware chunking: {len(documents)} chunks from {file_path}")
            return documents
            
        except ImportError:
            logger.warning("Chunkers not available, using RecursiveCharacterTextSplitter")
            # Legacy fallback
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
            )

            # Create documents from text
            chunks = splitter.create_documents([text])

            # Add metadata to each chunk
            for i, chunk in enumerate(chunks):
                chunk.metadata = {
                    "source": file_path,
                    "chunk_index": i,
                }

            logger.debug(
                f"Document chunked: {file_path} → {len(chunks)} chunks",
                extra={"file_path": file_path, "chunk_count": len(chunks)},
            )

            return chunks

    except Exception as e:
        logger.error(
            f"Failed to chunk document '{file_path}': {e}",
            extra={"file_path": file_path, "error": str(e)},
            exc_info=True,
        )
        raise


async def ingest_documents(
    file_paths: List[str],
    file_id: str,  # Required: Passed from API/Task
    embed_model: str = "nomic-embed-text",
    chunk_size: int = 500,
    chunk_overlap: int = 50,
    backend: str = "chroma",
) -> dict:
    """Ingest documents into vector store (Standardized).

    Args:
        file_paths: List of file paths to ingest
        file_id: Unique ID for the file (required)
        embed_model: Embedding model name
        chunk_size: Maximum characters per chunk
        chunk_overlap: Characters to overlap between chunks
        backend: Vector store backend ("chroma" or "qdrant")

    Returns:
        Dictionary with ingestion result
    """
    logger.info(
        f"Ingesting {len(file_paths)} documents",
        extra={
            "file_count": len(file_paths),
            "file_id": file_id,
            "embed_model": embed_model,
            "backend": backend,
        },
    )

    try:
        # Initialize vector store using Standard Abstraction
        from src.storage.chroma_store import ChromaVectorStore
        
        # Initialize store (always use standard collection)
        vector_store = ChromaVectorStore(
            collection_name=settings.CHROMA_COLLECTION,
            embed_model=embed_model
        )

        # Read all documents in parallel
        read_tasks = [read_document(fp) for fp in file_paths]
        contents = await asyncio.gather(*read_tasks, return_exceptions=True)

        all_chunks = []
        documents_processed = 0
        errors = []

        import uuid

        for file_path, content in zip(file_paths, contents):
            if isinstance(content, Exception):
                error_msg = f"Failed to read {file_path}: {str(content)}"
                logger.warning(error_msg)
                errors.append(error_msg)
                continue

            try:
                # Chunk the document
                chunks = chunk_document(
                    text=content,
                    file_path=file_path,
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                )
                
                # ENFORCE METADATA STANDARDS
                valid_chunks = []
                for i, chunk in enumerate(chunks):
                    # 1. Ensure absolute path for source
                    source_path = str(Path(file_path).resolve())
                    
                    # 2. Get chunk metadata or init new
                    meta = chunk.metadata if hasattr(chunk, "metadata") else {}
                    
                    # 3. Inject Critical Identity Fields
                    meta["file_id"] = file_id
                    meta["chunk_id"] = str(uuid.uuid4())
                    meta["source"] = source_path
                    meta["chunk_index"] = i
                    meta["total_chunks"] = len(chunks)
                    
                    # 4. Flatten metadata (ChromaDB requirement)
                    # Convert lists/dicts to JSON strings
                    import json
                    flat_meta = {}
                    for k, v in meta.items():
                        if isinstance(v, (list, dict)):
                            flat_meta[k] = json.dumps(v)
                        elif v is None:
                            flat_meta[k] = ""
                        else:
                            flat_meta[k] = v
                    
                    # 5. Create standardized chunk dict
                    valid_chunks.append({
                        "content": chunk.page_content,
                        "metadata": flat_meta
                    })
                
                all_chunks.extend(valid_chunks)
                documents_processed += 1
                
            except Exception as e:
                error_msg = f"Failed to chunk {file_path}: {str(e)}"
                logger.warning(error_msg)
                errors.append(error_msg)
                continue

        # Add to vector store
        if all_chunks:
            # Use standardized add_chunks method
            # Note: embeddings arg is empty list as Chroma generates them if not provided,
            # or we rely on vector_store implementation to handle embedding generation.
            # ChromaVectorStore uses vector_store.add_documents which handles embedding.
            # Wait, BaseVectorStore.add_chunks takes (chunks, embeddings).
            # But ChromaVectorStore implementation wraps `add_documents` which takes Document objects.
            # I should use `vector_store.add_chunks` if it generates embeddings?
            # BaseVectorStore interface says:
            # add_chunks(chunks: List[Dict], embeddings: List)
            # If embeddings is empty, does it generate? 
            # ChromaVectorStore implementation:
            # documents = [Document(...) for c in chunks]
            # self.client.add_documents(documents)
            # Chroma handles generation if embedding_function is set (which it is).
            
            await vector_store.add_chunks(chunks=all_chunks, embeddings=[])
            
            logger.info(
                f"Added {len(all_chunks)} chunks to vector store",
                extra={"chunk_count": len(all_chunks), "backend": backend},
            )

        result = {
            "success": documents_processed > 0,
            "documents_processed": documents_processed,
            "chunks_created": len(all_chunks),
            "backend": backend,
            "error": "; ".join(errors) if errors else None,
        }

        return result

    except Exception as e:
        logger.error(f"Document ingestion failed: {e}", exc_info=True)
        return {
            "success": False,
            "documents_processed": 0,
            "chunks_created": 0,
            "error": str(e),
        }


async def retrieve_docs(
    query: str,
    top_k: int = 5,
    embed_model: str = "nomic-embed-text",
    backend: str = "chroma",
    score_threshold: float = 0.5,
) -> List[dict]:
    """Retrieve relevant documents via semantic search (Delegated to Agent).

    Args:
        query: Search query string
        top_k: Number of results to return
        embed_model: Embedding model name
        backend: Vector store backend
        score_threshold: Minimum similarity score (0.0-1.0)

    Returns:
        List of dictionaries with id, content, metadata, score
    """
    logger.info(
        f"Retrieving documents (Delegated): query='{str(query)[:50]}...', top_k={top_k}",
        extra={"query_preview": str(query)[:50], "top_k": top_k},
    )

    try:
        # Phase 14: Delegate to RAGAgent (Single Source of Truth)
        # Prevent circular import
        from src.agents.rag.agent import get_shared_rag_agent
        
        agent = get_shared_rag_agent()
        
        # Use retrieve_with_reranking for best results
        # Note: We pass use_reranking=True explicitly as this is the "smart" tool
        result = await agent.retrieve_with_reranking(
            query=query,
            top_k=top_k,
            score_threshold=score_threshold,
            use_reranking=True
        )
        
        # Format results: retrieve_with_reranking returns a dict with "documents" list
        documents = result.get("documents", [])
        
        formatted_results = []
        for doc in documents:
            # Handle both ChunkResult objects and dicts
            if hasattr(doc, 'content'):
                formatted_results.append({
                    "id": getattr(doc, 'id', ''),
                    "content": getattr(doc, 'content', ''),
                    "metadata": getattr(doc, 'metadata', {}),
                    "score": getattr(doc, 'rerank_score', getattr(doc, 'score', 0.0))
                })
            else:
                formatted_results.append(doc)
                
        return formatted_results

    except Exception as e:
        logger.error(f"Document retrieval failed: {e}", exc_info=True)
        return []


async def generate_response(
    query: str,
    context: str,
    model_hint: Optional[str] = None,
) -> str:
    """Generate LLM response based on retrieved context.

    Uses ModelRouter to select optimal model for RAG tasks.

    Args:
        query: User query
        context: Retrieved context from documents
        model_hint: Optional model name hint (default: use ModelRouter selection)

    Returns:
        Generated response string

    Raises:
        Exception: If response generation fails
    """
    logger.info(
        f"Generating response for query: {str(query)[:50]}...",
        extra={"query_preview": str(query)[:50], "context_length": len(context)},
    )

    try:
        # Select model using ModelRouter
        if model_hint:
            model_name = model_hint
        else:
            model_name = model_router.select_model_by_task("rag_simple", prefer_local=False)  # Use cloud

        logger.debug(f"Using model: {model_name} for RAG response generation")

        # Build prompt template
        prompt = f"""You are a helpful assistant answering questions based on the provided context.

Context:
{context}

Question: {query}

Instructions:
- Answer based ONLY on the context provided
- If the context doesn't contain relevant information, say "I don't have enough information to answer that"
- Be concise and specific
- Cite relevant parts of the context when possible

Answer:"""

        # Invoke model with fallback
        chat_model = model_router.get_chat_model(model_name)
        response = await chat_model.ainvoke([{"role": "user", "content": prompt}])

        # Extract response content
        response_text = response.content if hasattr(response, "content") else str(response)

        logger.info(
            f"Response generated successfully ({len(response_text)} characters)",
            extra={"response_length": len(response_text), "model": model_name},
        )

        return response_text

    except Exception as e:
        logger.error(
            f"Response generation failed: {e}",
            extra={"error": str(e), "query_preview": str(query)[:50]},
            exc_info=True,
        )
        raise
